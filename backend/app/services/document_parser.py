"""
Extraction du texte d'un dossier de projet uploadé.

Remplace l'ancien décodage brut UTF-8 (inefficace sur les PDF compressés) par
une extraction réelle selon le type de fichier :
  - PDF  -> pypdf
  - DOCX -> python-docx (paragraphes + tableaux)
  - MD / TXT -> décodage UTF-8 direct
"""

import io


def _extract_pdf(contents: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(contents))
    parts = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text:
            parts.append(text)
    return "\n".join(parts)


def _extract_docx(contents: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(contents))
    parts = [p.text for p in document.paragraphs if p.text]

    # Inclure aussi le texte contenu dans les tableaux (fréquent dans un dossier).
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells if cell.text]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_plain(contents: bytes) -> str:
    return contents.decode("utf-8", errors="ignore")


def extract_text(filename: str, contents: bytes) -> str:
    """
    Extrait le texte d'un fichier à partir de son nom (pour le type) et de son
    contenu binaire. Retourne une chaîne vide si l'extraction échoue.
    """
    ext = ""
    if filename and "." in filename:
        ext = "." + filename.rsplit(".", 1)[-1].lower()

    try:
        if ext == ".pdf":
            return _extract_pdf(contents)
        if ext == ".docx":
            return _extract_docx(contents)
        # .md, .txt et tout autre format texte
        return _extract_plain(contents)
    except Exception as e:
        print(f"[document_parser] Extraction échouée pour '{filename}' ({ext}): {e}")
        return ""
